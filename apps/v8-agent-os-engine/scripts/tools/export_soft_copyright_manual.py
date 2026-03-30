from __future__ import annotations

import argparse
import json
import math
import re
from pathlib import Path

from docx import Document
from soft_copyright_common import (
    apply_a4_vertical_section,
    configure_manual_styles,
    load_bundle_meta,
    set_footer_empty,
    set_header,
    set_run_font,
)

NUMBERED_RE = re.compile(r"^\d+\.\s+")


def _count_visual_lines(text: str, width: int = 36) -> int:
    if not text:
        return 1
    return max(1, math.ceil(len(text) / width))


def _table_row_to_paragraph(line: str) -> str:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return "  ".join(cell for cell in cells if cell and not set(cell) <= {"-", ":"})


def _apply_runs(paragraph, *, font_name: str, size_pt: float, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, name=font_name, size_pt=size_pt, bold=bold)


def _add_paragraph(document: Document, line: str, stats: dict[str, int]) -> None:
    stripped = line.rstrip()
    if not stripped:
        return
    if stripped.startswith("# "):
        paragraph = document.add_heading(stripped[2:].strip(), level=1)
        _apply_runs(paragraph, font_name="黑体", size_pt=14, bold=True)
        stats["headingCount"] += 1
        stats["estimatedVisualLines"] += 1
        return
    if stripped.startswith("## "):
        paragraph = document.add_heading(stripped[3:].strip(), level=2)
        _apply_runs(paragraph, font_name="黑体", size_pt=12, bold=True)
        stats["headingCount"] += 1
        stats["estimatedVisualLines"] += 1
        return
    if stripped.startswith("### "):
        paragraph = document.add_heading(stripped[4:].strip(), level=3)
        _apply_runs(paragraph, font_name="黑体", size_pt=12, bold=True)
        stats["headingCount"] += 1
        stats["estimatedVisualLines"] += 1
        return
    if stripped.startswith("- "):
        text = stripped[2:].strip()
        paragraph = document.add_paragraph(text)
        _apply_runs(paragraph, font_name="宋体", size_pt=12)
        stats["paragraphCount"] += 1
        stats["estimatedVisualLines"] += _count_visual_lines(text)
        return
    if NUMBERED_RE.match(stripped):
        text = NUMBERED_RE.sub("", stripped, count=1).strip()
        paragraph = document.add_paragraph(text)
        _apply_runs(paragraph, font_name="宋体", size_pt=12)
        stats["paragraphCount"] += 1
        stats["estimatedVisualLines"] += _count_visual_lines(text)
        return
    if stripped.startswith("|") and stripped.endswith("|"):
        text = _table_row_to_paragraph(stripped)
        if text:
            paragraph = document.add_paragraph(text)
            _apply_runs(paragraph, font_name="宋体", size_pt=12)
            stats["paragraphCount"] += 1
            stats["estimatedVisualLines"] += _count_visual_lines(text)
        return
    paragraph = document.add_paragraph(stripped)
    _apply_runs(paragraph, font_name="宋体", size_pt=12)
    stats["paragraphCount"] += 1
    stats["estimatedVisualLines"] += _count_visual_lines(stripped)


def export_manual(repo_root: Path, source: Path, output: Path, stats_output: Path) -> None:
    text = source.read_text(encoding="utf-8")
    meta = load_bundle_meta(repo_root)
    document = Document()
    section = document.sections[0]
    apply_a4_vertical_section(section)
    configure_manual_styles(document)
    set_header(section, f"{meta['softwareName']} {meta['version']} {meta['manualHeaderLabel']}")
    set_footer_empty(section)

    stats = {
        "softwareName": meta["softwareName"],
        "version": meta["version"],
        "documentName": meta["documentName"],
        "paragraphCount": 0,
        "headingCount": 0,
        "estimatedVisualLines": 0,
    }
    for line in text.splitlines():
        _add_paragraph(document, line, stats)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    stats["estimatedPageCount"] = max(1, math.ceil(stats["estimatedVisualLines"] / 30))
    stats["headerLeftText"] = f"{meta['softwareName']} {meta['version']} {meta['manualHeaderLabel']}"
    stats["headerRightText"] = "第X页 / 共Y页（Word 域）"
    stats_output.write_text(json.dumps(stats, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="Export software copyright manual markdown to DOCX.")
    parser.add_argument("--repo-root", required=True, type=Path)
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--stats-output", type=Path)
    args = parser.parse_args()
    stats_output = args.stats_output or args.output.with_name(f"{args.output.stem}_stats.json")
    export_manual(args.repo_root, args.source, args.output, stats_output)


if __name__ == "__main__":
    main()
