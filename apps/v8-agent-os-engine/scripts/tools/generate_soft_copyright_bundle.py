from __future__ import annotations

import argparse
import json
import math
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from soft_copyright_common import (
    apply_a4_vertical_section,
    configure_source_styles,
    load_bundle_meta,
    set_footer_empty,
    set_header,
    set_run_font,
)

EXCERPT_TARGET_LINES = 3000
LINES_PER_PAGE = 50
FULL_TARGET_LINES = 3200

ENGINE_TARGET_LINES = 2600
ADMIN_TARGET_LINES = 300
WEB_TARGET_LINES = 300
MIN_FILE_LINES = 80
MIN_ENGINE_FILES = 6
MIN_ADMIN_FILES = 3
MIN_WEB_FILES = 3
MAX_SOURCE_LINE_LENGTH = 70

PREFERRED_ENGINE_FILES = [
    "apps/v8-agent-os-engine/main.py",
    "apps/v8-agent-os-engine/erc/kernel.py",
    "apps/v8-agent-os-engine/erc/command_service.py",
    "apps/v8-agent-os-engine/erc/command_router.py",
    "apps/v8-agent-os-engine/erc/capability_registry.py",
    "apps/v8-agent-os-engine/api/chat_realtime_routes.py",
    "apps/v8-agent-os-engine/api/network_supervisor_routes.py",
    "apps/v8-agent-os-engine/runtimes/network_supervisor/service.py",
    "apps/v8-agent-os-engine/runtimes/network_supervisor/runtime.py",
    "apps/v8-agent-os-engine/core/extensions_runtime.py",
    "apps/v8-agent-os-engine/core/storage.py",
]

PREFERRED_ADMIN_FILES = [
    "apps/v8-agent-os-admin/src/components/network-supervisor/NetworkSupervisorRuntimeWorkbench.tsx",
    "apps/v8-agent-os-admin/src/components/runtime/RuntimeGovernanceWorkbench.tsx",
    "apps/v8-agent-os-admin/src/components/plugin-host/PluginHostWorkbench.tsx",
    "apps/v8-agent-os-admin/src/components/memory/GraphViewer.tsx",
    "apps/v8-agent-os-admin/src/components/rpa/RPAWorkbench.tsx",
]

PREFERRED_WEB_FILES = [
    "apps/v8-agent-os-web/src/app/chat/ChatClient.tsx",
    "apps/v8-agent-os-web/src/components/chat/InputArea.tsx",
    "apps/v8-agent-os-web/src/lib/chat-stream-state.ts",
    "apps/v8-agent-os-web/src/hooks/use-langgraph-stream.ts",
    "apps/v8-agent-os-web/src/components/chat/ChatMessage.tsx",
    "apps/v8-agent-os-web/src/components/chat/RuntimeTimelinePanel.tsx",
]

EXCLUDE_PARTS = {
    "node_modules",
    ".next",
    ".venv",
    "venv",
    "__pycache__",
    "dist",
    "build",
    ".pytest_cache",
    ".mypy_cache",
    ".ruff_cache",
    ".turbo",
    ".git",
}

INLINE_PATH_PREFIX_RE = re.compile(r"^[A-Za-z0-9_.-]+(?:/[A-Za-z0-9_.-]+)+:\d{4}\s+\|\s+")


@dataclass
class SourceLine:
    domain: str
    relative_path: str
    line_number: int
    content: str


@dataclass
class FilterStats:
    blank_removed: int = 0
    comment_removed: int = 0
    long_removed: int = 0


def _is_excluded(path: Path) -> bool:
    return any(part in EXCLUDE_PARTS for part in path.parts)


def _priority_score(path: Path, rules: list[str]) -> tuple[int, str]:
    normalized = str(path).replace("\\", "/")
    for index, rule in enumerate(rules):
        if normalized.endswith(rule) or f"/{rule}" in normalized:
            return (index, normalized)
    return (len(rules), normalized)


def _collect_files(root: Path, suffixes: tuple[str, ...], rules: list[str]) -> list[Path]:
    files: list[Path] = []
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        if _is_excluded(path):
            continue
        if path.suffix.lower() not in suffixes:
            continue
        files.append(path)
    return sorted(files, key=lambda item: _priority_score(item, rules))


def _strip_inline_path_prefix(line: str) -> str:
    return INLINE_PATH_PREFIX_RE.sub("", line)


def _iter_effective_lines(path: Path, repo_root: Path, domain: str, stats: FilterStats) -> list[SourceLine]:
    relative_path = path.relative_to(repo_root).as_posix()
    language = path.suffix.lower()
    raw_lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()
    result: list[SourceLine] = []
    in_block_comment = False
    block_end = ""

    for idx, raw_line in enumerate(raw_lines, start=1):
        line = _strip_inline_path_prefix(raw_line.rstrip("\n"))
        stripped = line.strip()

        if in_block_comment:
            stats.comment_removed += 1
            if block_end and block_end in stripped:
                in_block_comment = False
            continue

        if not stripped:
            stats.blank_removed += 1
            continue

        if language == ".py":
            if stripped.startswith("#"):
                stats.comment_removed += 1
                continue
            if stripped.startswith(('"""', "'''")):
                stats.comment_removed += 1
                if not ((stripped.endswith('"""') and len(stripped) > 3) or (stripped.endswith("'''") and len(stripped) > 3)):
                    in_block_comment = True
                    block_end = stripped[:3]
                continue
        else:
            if stripped.startswith("//"):
                stats.comment_removed += 1
                continue
            if stripped.startswith("/*"):
                stats.comment_removed += 1
                if "*/" not in stripped[2:]:
                    in_block_comment = True
                    block_end = "*/"
                continue
            if stripped.startswith("*") or stripped.startswith("*/"):
                stats.comment_removed += 1
                continue

        if len(line.expandtabs(4)) > MAX_SOURCE_LINE_LENGTH:
            stats.long_removed += 1
            continue

        result.append(
            SourceLine(
                domain=domain,
                relative_path=relative_path,
                line_number=idx,
                content=line,
            )
        )
    return result


def _count_effective_lines(path: Path, repo_root: Path, domain: str) -> int:
    stats = FilterStats()
    return len(_iter_effective_lines(path, repo_root, domain, stats))


def _select_files(
    repo_root: Path,
    root: Path,
    domain: str,
    preferred: list[str],
    rules: list[str],
    suffixes: tuple[str, ...],
    target_lines: int,
    min_files: int,
) -> list[Path]:
    selected: list[Path] = []
    selected_set: set[Path] = set()
    total = 0

    for relative in preferred:
        path = repo_root / Path(relative)
        if not path.exists() or not path.is_file():
            continue
        effective_lines = _count_effective_lines(path, repo_root, domain)
        if effective_lines < MIN_FILE_LINES:
            continue
        selected.append(path)
        selected_set.add(path)
        total += effective_lines
        if total >= target_lines and len(selected) >= min_files:
            break

    if total >= target_lines and len(selected) >= min_files:
        return selected

    for path in _collect_files(root, suffixes, rules):
        if path in selected_set:
            continue
        effective_lines = _count_effective_lines(path, repo_root, domain)
        if effective_lines < MIN_FILE_LINES:
            continue
        selected.append(path)
        selected_set.add(path)
        total += effective_lines
        if total >= target_lines and len(selected) >= min_files:
            return selected

    raise RuntimeError(f"{domain} 域可用有效代码不足：累计 {total} 行，文件数 {len(selected)}。")


def _take_balanced_source_lines(
    domain: str,
    repo_root: Path,
    files: list[Path],
    target_lines: int,
    min_files: int,
    aggregate_stats: FilterStats,
) -> list[SourceLine]:
    if len(files) < min_files:
        raise RuntimeError(f"{domain} 域选定文件不足最小文件数要求。")

    buffers: list[tuple[Path, list[SourceLine]]] = []
    for path in files:
        local_stats = FilterStats()
        lines = _iter_effective_lines(path, repo_root, domain, local_stats)
        aggregate_stats.blank_removed += local_stats.blank_removed
        aggregate_stats.comment_removed += local_stats.comment_removed
        buffers.append((path, lines))

    allotment = max(1, target_lines // min_files)
    selected: list[SourceLine] = []
    taken: dict[Path, int] = {}
    remaining = target_lines

    for path, lines in buffers:
        take = min(len(lines), allotment, remaining)
        selected.extend(lines[:take])
        taken[path] = take
        remaining -= take
        if remaining == 0:
            break

    if remaining > 0:
        for path, lines in buffers:
            start = taken.get(path, 0)
            extra = min(len(lines) - start, remaining)
            selected.extend(lines[start : start + extra])
            taken[path] = start + extra
            remaining -= extra
            if remaining == 0:
                break

    if remaining > 0:
        raise RuntimeError(f"{domain} 域有效代码不足 {target_lines} 行，仍缺少 {remaining} 行。")

    return selected


def _write_text(path: Path, lines: list[str]) -> None:
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _write_docx(path: Path, header_text: str, lines: list[str]) -> None:
    document = Document()
    section = document.sections[0]
    apply_a4_vertical_section(section)
    configure_source_styles(document)
    set_header(section, header_text, total_pages=len(lines) // LINES_PER_PAGE)
    set_footer_empty(section)

    total_pages = len(lines) // LINES_PER_PAGE
    for page_index in range(total_pages):
        chunk = lines[page_index * LINES_PER_PAGE : (page_index + 1) * LINES_PER_PAGE]
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.left_indent = 0
        paragraph.paragraph_format.right_indent = 0
        paragraph.paragraph_format.first_line_indent = 0
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(12)
        run = paragraph.add_run()
        set_run_font(run, name="宋体", size_pt=12)
        for line_index, line in enumerate(chunk):
            run.add_text(line)
            if line_index < len(chunk) - 1:
                run.add_break(WD_BREAK.LINE)
        if page_index < total_pages - 1:
            paragraph.add_run().add_break(WD_BREAK.PAGE)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _build_manifest(entries: list[SourceLine]) -> dict:
    grouped: dict[tuple[str, str], dict] = {}
    for item in entries:
        key = (item.domain, item.relative_path)
        current = grouped.setdefault(
            key,
            {
                "domain": item.domain,
                "relativePath": item.relative_path,
                "usedLineCount": 0,
                "startLine": item.line_number,
                "endLine": item.line_number,
            },
        )
        current["usedLineCount"] += 1
        current["startLine"] = min(current["startLine"], item.line_number)
        current["endLine"] = max(current["endLine"], item.line_number)
    return {"totalFiles": len(grouped), "entries": list(grouped.values())}


def generate(repo_root: Path, output_dir: Path) -> None:
    meta = load_bundle_meta(repo_root)

    engine_root = repo_root / "apps" / "v8-agent-os-engine"
    admin_root = repo_root / "apps" / "v8-agent-os-admin" / "src"
    web_root = repo_root / "apps" / "v8-agent-os-web" / "src"

    engine_rules = [
        "main.py",
        "erc/kernel.py",
        "erc/runtime_registry.py",
        "erc/command_service.py",
        "erc/command_router.py",
        "erc/capability_registry.py",
        "api/chat_realtime_routes.py",
        "api/network_supervisor_routes.py",
        "runtimes/network_supervisor/service.py",
        "runtimes/network_supervisor/runtime.py",
        "runtimes/chat/runtime.py",
        "runtimes/memory/runtime.py",
        "core/storage.py",
        "core/extensions_runtime.py",
        "erc/",
        "runtimes/",
        "api/",
        "core/",
        "graph/",
        "skills/",
        "persistence/repositories/",
    ]
    admin_rules = [
        "components/network-supervisor/",
        "components/runtime/",
        "components/plugin-host/",
        "components/memory/",
        "components/rpa/",
        "components/automation/",
        "components/",
        "lib/",
        "app/admin/",
        "app/api/",
        "app/layout.tsx",
        "app/page.tsx",
    ]
    web_rules = [
        "components/chat/",
        "store/",
        "hooks/",
        "components/connection/",
        "components/layout/",
        "lib/",
        "app/chat/",
        "app/api/",
        "app/connect/",
        "app/layout.tsx",
        "app/page.tsx",
    ]

    selected_engine_files = _select_files(
        repo_root,
        engine_root,
        "engine",
        PREFERRED_ENGINE_FILES,
        engine_rules,
        (".py",),
        ENGINE_TARGET_LINES,
        MIN_ENGINE_FILES,
    )
    selected_admin_files = _select_files(
        repo_root,
        admin_root,
        "admin",
        PREFERRED_ADMIN_FILES,
        admin_rules,
        (".ts", ".tsx", ".js", ".jsx"),
        ADMIN_TARGET_LINES,
        MIN_ADMIN_FILES,
    )
    selected_web_files = _select_files(
        repo_root,
        web_root,
        "web",
        PREFERRED_WEB_FILES,
        web_rules,
        (".ts", ".tsx", ".js", ".jsx"),
        WEB_TARGET_LINES,
        MIN_WEB_FILES,
    )

    filter_stats = FilterStats()
    engine_lines = _take_balanced_source_lines("engine", repo_root, selected_engine_files, ENGINE_TARGET_LINES, MIN_ENGINE_FILES, filter_stats)
    admin_lines = _take_balanced_source_lines("admin", repo_root, selected_admin_files, ADMIN_TARGET_LINES, MIN_ADMIN_FILES, filter_stats)
    web_lines = _take_balanced_source_lines("web", repo_root, selected_web_files, WEB_TARGET_LINES, MIN_WEB_FILES, filter_stats)

    full_lines = engine_lines + admin_lines + web_lines
    if len(full_lines) < FULL_TARGET_LINES:
        raise RuntimeError(f"总有效代码不足 {FULL_TARGET_LINES} 行，当前仅有 {len(full_lines)} 行。")

    excerpt_lines = full_lines[: EXCERPT_TARGET_LINES // 2] + full_lines[-EXCERPT_TARGET_LINES // 2 :]
    if len(excerpt_lines) != EXCERPT_TARGET_LINES:
        raise RuntimeError("节选源码行数不等于 3000 行。")

    manifest = _build_manifest(full_lines)
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "source_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    manifest_md = ["# 申报源程序清单", ""]
    for domain in ("engine", "admin", "web"):
        manifest_md.append(f"## {domain}")
        for item in manifest["entries"]:
            if item["domain"] != domain:
                continue
            manifest_md.append(
                f"- `{item['relativePath']}`（使用 {item['usedLineCount']} 行，源行号 {item['startLine']}-{item['endLine']}）"
            )
        manifest_md.append("")
    _write_text(output_dir / "source_manifest.md", manifest_md)

    full_text_lines = [item.content for item in full_lines]
    excerpt_text_lines = [item.content for item in excerpt_lines]

    _write_text(output_dir / "source_program_full.txt", full_text_lines)
    _write_text(output_dir / "source_program_excerpt.txt", excerpt_text_lines)
    _write_docx(
        output_dir / "source_program_excerpt.docx",
        f"{meta['softwareName']} {meta['version']} {meta['sourceHeaderLabel']}",
        excerpt_text_lines,
    )

    stats = {
        "softwareName": meta["softwareName"],
        "version": meta["version"],
        "headerText": f"{meta['softwareName']} {meta['version']} {meta['sourceHeaderLabel']}",
        "fullLineCount": len(full_lines),
        "excerptLineCount": len(excerpt_lines),
        "effectiveCodeLines": len(excerpt_lines),
        "commentLinesRemoved": filter_stats.comment_removed,
        "blankLinesRemoved": filter_stats.blank_removed,
        "longLinesRemoved": filter_stats.long_removed,
        "linesPerPage": LINES_PER_PAGE,
        "pageCount": len(excerpt_lines) // LINES_PER_PAGE,
        "targetPages": 60,
        "selectedFilesByDomain": {
            "engine": len(selected_engine_files),
            "admin": len(selected_admin_files),
            "web": len(selected_web_files),
        },
        "selectedFiles": [item["relativePath"] for item in manifest["entries"]],
    }
    (output_dir / "source_program_stats.json").write_text(json.dumps(stats, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate software copyright source-program bundle.")
    parser.add_argument("--repo-root", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    args = parser.parse_args()
    generate(args.repo_root, args.output_dir)


if __name__ == "__main__":
    main()
