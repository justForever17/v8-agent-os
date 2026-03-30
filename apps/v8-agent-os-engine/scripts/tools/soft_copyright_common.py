from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SOFTWARE_COPYRIGHT_DIR = Path("docs") / "software-copyright"
META_PATH = SOFTWARE_COPYRIGHT_DIR / "bundle_meta.json"
A4_WIDTH_CM = 21.0
A4_HEIGHT_CM = 29.7
PAGE_MARGIN_CM = 2.5
CONTENT_WIDTH_CM = A4_WIDTH_CM - PAGE_MARGIN_CM * 2


def load_bundle_meta(repo_root: Path) -> dict[str, Any]:
    meta_path = repo_root / META_PATH
    return json.loads(meta_path.read_text(encoding="utf-8"))


def apply_a4_vertical_section(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(A4_WIDTH_CM)
    section.page_height = Cm(A4_HEIGHT_CM)
    section.top_margin = Cm(PAGE_MARGIN_CM)
    section.bottom_margin = Cm(PAGE_MARGIN_CM)
    section.left_margin = Cm(PAGE_MARGIN_CM)
    section.right_margin = Cm(PAGE_MARGIN_CM)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.gutter = Cm(0)


def _set_rfonts(r_pr, font_name: str) -> None:
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attr), font_name)


def _set_style_font(style, font_name: str, size_pt: float, *, bold: bool = False) -> None:
    style.font.name = font_name
    _set_rfonts(style._element.get_or_add_rPr(), font_name)
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def set_run_font(run, *, name: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = name
    _set_rfonts(run._element.get_or_add_rPr(), name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.highlight_color = None


def add_field_run(paragraph, field_name: str, *, font_name: str = "宋体", size_pt: float = 12) -> None:
    run = paragraph.add_run()
    set_run_font(run, name=font_name, size_pt=size_pt)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    run._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    text_node = OxmlElement("w:t")
    text_node.text = "1"
    run._r.append(text_node)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def _clear_header_footer(container) -> None:
    element = container._element
    for child in list(element):
        element.remove(child)


def _set_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    else:
        for child in list(tbl_borders):
            tbl_borders.remove(child)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        tbl_borders.append(border)


def set_header(section, left_text: str, total_pages: int | None = None) -> None:
    header = section.header
    _clear_header_footer(header)

    table = header.add_table(rows=1, cols=2, width=Cm(CONTENT_WIDTH_CM))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders_none(table)

    left_width = Cm(CONTENT_WIDTH_CM * 0.62)
    right_width = Cm(CONTENT_WIDTH_CM * 0.38)
    table.columns[0].width = left_width
    table.columns[1].width = right_width

    left_cell, right_cell = table.rows[0].cells
    left_cell.width = left_width
    right_cell.width = right_width

    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_paragraph.paragraph_format.space_before = Pt(0)
    left_paragraph.paragraph_format.space_after = Pt(0)
    left_run = left_paragraph.add_run(left_text)
    set_run_font(left_run, name="宋体", size_pt=12)

    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_paragraph.paragraph_format.space_before = Pt(0)
    right_paragraph.paragraph_format.space_after = Pt(0)
    label_run = right_paragraph.add_run("第")
    set_run_font(label_run, name="宋体", size_pt=12)
    add_field_run(right_paragraph, "PAGE", font_name="宋体", size_pt=12)
    middle_run = right_paragraph.add_run("页 / 共")
    set_run_font(middle_run, name="宋体", size_pt=12)
    if total_pages is None:
        add_field_run(right_paragraph, "NUMPAGES", font_name="宋体", size_pt=12)
    else:
        pages_run = right_paragraph.add_run(str(total_pages))
        set_run_font(pages_run, name="宋体", size_pt=12)
    tail_run = right_paragraph.add_run("页")
    set_run_font(tail_run, name="宋体", size_pt=12)


def set_footer_empty(section) -> None:
    footer = section.footer
    _clear_header_footer(footer)
    paragraph = footer.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def configure_source_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    _set_style_font(normal, "宋体", 12)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.left_indent = Pt(0)
    normal.paragraph_format.right_indent = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(12)


def configure_manual_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    _set_style_font(normal, "宋体", 12)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.left_indent = Pt(0)
    normal.paragraph_format.right_indent = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(20)

    for style_name, font_name, size_pt, bold in (
        ("Heading 1", "黑体", 14, True),
        ("Heading 2", "黑体", 12, True),
        ("Heading 3", "黑体", 12, True),
    ):
        style = document.styles[style_name]
        _set_style_font(style, font_name, size_pt, bold=bold)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.left_indent = Pt(0)
        style.paragraph_format.right_indent = Pt(0)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(20)


def is_comment_or_blank(line: str, *, language_hint: str = "") -> bool:
    stripped = line.strip()
    if not stripped:
        return True
    if stripped.startswith(("#", "//", "/*", "*/")):
        return True
    if stripped == "*" or stripped.startswith("* "):
        return True
    if stripped in {'"""', "'''"}:
        return True
    if stripped.startswith('"""') or stripped.startswith("'''"):
        if language_hint == "python" or stripped.endswith('"""') or stripped.endswith("'''"):
            return True
    return False
